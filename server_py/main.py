import os
import json
import sqlite3
import uuid
from datetime import datetime
from typing import Optional, List, Dict, Any
from fastapi import FastAPI, HTTPException, UploadFile, File, Response, Query
from fastapi.responses import FileResponse
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel

# Initialize FastAPI App
app = FastAPI(title="EduForge Python API", version="2.0.0")

# Enable CORS for frontend Vite & local tools
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

DB_PATH = os.path.join(os.path.dirname(__file__), "eduforge.db")
UPLOADS_DIR = os.path.join(os.path.dirname(__file__), "uploads")
os.makedirs(UPLOADS_DIR, exist_ok=True)


def get_db():
    conn = sqlite3.connect(DB_PATH)
    conn.row_factory = sqlite3.Row
    return conn


def init_db():
    conn = get_db()
    cursor = conn.cursor()
    
    # Documents table
    cursor.execute("""
    CREATE TABLE IF NOT EXISTS documents (
        id TEXT PRIMARY KEY,
        title TEXT NOT NULL,
        template_id TEXT,
        settings TEXT NOT NULL,
        metadata TEXT NOT NULL,
        sections TEXT NOT NULL,
        created_at TEXT NOT NULL,
        updated_at TEXT NOT NULL
    )
    """)

    # Question Bank table
    cursor.execute("""
    CREATE TABLE IF NOT EXISTS question_bank (
        id TEXT PRIMARY KEY,
        question_number INTEGER,
        question_type TEXT NOT NULL,
        content TEXT NOT NULL,
        raw_text TEXT,
        options TEXT NOT NULL,
        correct_answer TEXT,
        marks REAL NOT NULL,
        negative_marks REAL DEFAULT 0,
        subject TEXT,
        chapter TEXT,
        topic TEXT,
        difficulty TEXT NOT NULL,
        tags TEXT NOT NULL,
        year TEXT,
        option_layout TEXT NOT NULL,
        explanation_text TEXT,
        diagram_svg TEXT,
        diagram_url TEXT,
        is_system INTEGER DEFAULT 0,
        created_at TEXT NOT NULL,
        updated_at TEXT NOT NULL
    )
    """)

    # Templates table
    cursor.execute("""
    CREATE TABLE IF NOT EXISTS templates (
        id TEXT PRIMARY KEY,
        name TEXT NOT NULL,
        description TEXT,
        category TEXT NOT NULL,
        settings TEXT NOT NULL,
        default_metadata TEXT NOT NULL,
        default_sections TEXT NOT NULL,
        created_at TEXT NOT NULL
    )
    """)

    # Safe migrations for new columns
    try:
        cursor.execute("ALTER TABLE question_bank ADD COLUMN diagram_svg TEXT")
    except sqlite3.OperationalError:
        pass

    try:
        cursor.execute("ALTER TABLE question_bank ADD COLUMN diagram_url TEXT")
    except sqlite3.OperationalError:
        pass

    # Check question count and seed with full set if < 10
    cursor.execute("SELECT COUNT(*) as count FROM question_bank")
    if cursor.fetchone()["count"] < 10:
        cursor.execute("DELETE FROM question_bank")
        seed_questions(cursor)

    # Seed Templates if empty
    cursor.execute("SELECT COUNT(*) as count FROM templates")
    if cursor.fetchone()["count"] == 0:
        seed_templates(cursor)

    conn.commit()
    conn.close()


def seed_questions(cursor):
    sample_questions = [
        {
            "id": "q-phys-001",
            "question_number": 1,
            "question_type": "MCQ_SINGLE",
            "content": json.dumps([]),
            "raw_text": "A particle is projected from the ground with an initial velocity u at an angle \\theta with the horizontal. The maximum height H reached by the projectile is given by:",
            "options": json.dumps([
                {"id": "opt-1a", "key": "a", "rawText": "H = \\frac{u^2 \\sin^2\\theta}{2g}", "isCorrect": True, "content": []},
                {"id": "opt-1b", "key": "b", "rawText": "H = \\frac{u^2 \\sin 2\\theta}{g}", "isCorrect": False, "content": []},
                {"id": "opt-1c", "key": "c", "rawText": "H = \\frac{u^2 \\cos^2\\theta}{2g}", "isCorrect": False, "content": []},
                {"id": "opt-1d", "key": "d", "rawText": "H = \\frac{2u \\sin\\theta}{g}", "isCorrect": False, "content": []}
            ]),
            "correct_answer": "a",
            "marks": 4,
            "negative_marks": 1,
            "subject": "Physics",
            "chapter": "Kinematics",
            "topic": "Projectile Motion",
            "difficulty": "Medium",
            "tags": json.dumps(["kinematics", "projectile", "mechanics"]),
            "year": "2025",
            "option_layout": "grid_2x2",
            "explanation_text": "At maximum height v_y = 0. Using v_y^2 = u_y^2 - 2gH => H = (u^2 sin^2 theta)/(2g).",
            "diagram_svg": None,
            "diagram_url": None,
            "is_system": 1,
            "created_at": datetime.utcnow().isoformat(),
            "updated_at": datetime.utcnow().isoformat()
        },
        {
            "id": "q-phys-002",
            "question_number": 2,
            "question_type": "MCQ_SINGLE",
            "content": json.dumps([]),
            "raw_text": "An electric dipole of dipole moment \\vec{p} is placed in a uniform electric field \\vec{E}. The torque acting on the dipole is:",
            "options": json.dumps([
                {"id": "opt-2a", "key": "a", "rawText": "\\vec{\\tau} = \\vec{p} \\times \\vec{E}", "isCorrect": True, "content": []},
                {"id": "opt-2b", "key": "b", "rawText": "\\vec{\\tau} = \\vec{p} \\cdot \\vec{E}", "isCorrect": False, "content": []},
                {"id": "opt-2c", "key": "c", "rawText": "\\vec{\\tau} = \\vec{E} \\times \\vec{p}", "isCorrect": False, "content": []},
                {"id": "opt-2d", "key": "d", "rawText": "\\vec{\\tau} = \\vec{0}", "isCorrect": False, "content": []}
            ]),
            "correct_answer": "a",
            "marks": 4,
            "negative_marks": 1,
            "subject": "Physics",
            "chapter": "Electrostatics",
            "topic": "Electric Dipole",
            "difficulty": "Easy",
            "tags": json.dumps(["electrostatics", "dipole", "torque"]),
            "year": "2025",
            "option_layout": "grid_2x2",
            "explanation_text": "Torque tau = r x F = q(2a) x E = p x E.",
            "diagram_svg": None,
            "diagram_url": None,
            "is_system": 1,
            "created_at": datetime.utcnow().isoformat(),
            "updated_at": datetime.utcnow().isoformat()
        },
        {
            "id": "q-phys-003",
            "question_number": 3,
            "question_type": "MCQ_SINGLE",
            "content": json.dumps([]),
            "raw_text": "In a series L-C-R alternating current circuit with inductance L, capacitance C, and resistance R, electrical resonance occurs at angular frequency:",
            "options": json.dumps([
                {"id": "opt-3a", "key": "a", "rawText": "\\omega_0 = \\frac{1}{\\sqrt{LC}}", "isCorrect": True, "content": []},
                {"id": "opt-3b", "key": "b", "rawText": "\\omega_0 = \\sqrt{LC}", "isCorrect": False, "content": []},
                {"id": "opt-3c", "key": "c", "rawText": "\\omega_0 = \\frac{L}{C}", "isCorrect": False, "content": []},
                {"id": "opt-3d", "key": "d", "rawText": "\\omega_0 = \\frac{1}{LC}", "isCorrect": False, "content": []}
            ]),
            "correct_answer": "a",
            "marks": 4,
            "negative_marks": 1,
            "subject": "Physics",
            "chapter": "Alternating Current",
            "topic": "Resonance",
            "difficulty": "Easy",
            "tags": json.dumps(["ac-circuits", "resonance", "lcr"]),
            "year": "2024",
            "option_layout": "grid_2x2",
            "explanation_text": "At resonance X_L = X_C => omega L = 1/(omega C) => omega = 1/sqrt(LC).",
            "diagram_svg": None,
            "diagram_url": None,
            "is_system": 1,
            "created_at": datetime.utcnow().isoformat(),
            "updated_at": datetime.utcnow().isoformat()
        },
        {
            "id": "q-phys-004",
            "question_number": 4,
            "question_type": "MCQ_SINGLE",
            "content": json.dumps([]),
            "raw_text": "Monochromatic light of frequency \\nu is incident on a photosensitive metal plate of work function \\Phi. The maximum kinetic energy of emitted photoelectrons is:",
            "options": json.dumps([
                {"id": "opt-4a", "key": "a", "rawText": "K_{\\max} = h\\nu - \\Phi", "isCorrect": True, "content": []},
                {"id": "opt-4b", "key": "b", "rawText": "K_{\\max} = h\\nu + \\Phi", "isCorrect": False, "content": []},
                {"id": "opt-4c", "key": "c", "rawText": "K_{\\max} = \\frac{h\\nu}{\\Phi}", "isCorrect": False, "content": []},
                {"id": "opt-4d", "key": "d", "rawText": "K_{\\max} = \\Phi - h\\nu", "isCorrect": False, "content": []}
            ]),
            "correct_answer": "a",
            "marks": 4,
            "negative_marks": 1,
            "subject": "Physics",
            "chapter": "Dual Nature of Radiation",
            "topic": "Photoelectric Effect",
            "difficulty": "Easy",
            "tags": json.dumps(["photoelectric", "quantum", "modern-physics"]),
            "year": "2025",
            "option_layout": "grid_2x2",
            "explanation_text": "Einstein photoelectric equation: E = h nu = Phi + K_max.",
            "diagram_svg": None,
            "diagram_url": None,
            "is_system": 1,
            "created_at": datetime.utcnow().isoformat(),
            "updated_at": datetime.utcnow().isoformat()
        },
        {
            "id": "q-chem-005",
            "question_number": 5,
            "question_type": "MCQ_SINGLE",
            "content": json.dumps([]),
            "raw_text": "For the exothermic industrial Haber synthesis: N₂(g) + 3H₂(g) ⇌ 2NH₃(g), ΔH = -92.4 kJ/mol, which conditions maximize ammonia yield?",
            "options": json.dumps([
                {"id": "opt-5a", "key": "a", "rawText": "High Pressure (~200 atm) and Optimum Temperature (~450°C)", "isCorrect": True, "content": []},
                {"id": "opt-5b", "key": "b", "rawText": "Low Pressure and High Temperature (>900°C)", "isCorrect": False, "content": []},
                {"id": "opt-5c", "key": "c", "rawText": "Low Pressure and Low Temperature", "isCorrect": False, "content": []},
                {"id": "opt-5d", "key": "d", "rawText": "High Pressure and Room Temperature without catalyst", "isCorrect": False, "content": []}
            ]),
            "correct_answer": "a",
            "marks": 4,
            "negative_marks": 1,
            "subject": "Chemistry",
            "chapter": "Chemical Equilibrium",
            "topic": "Haber Process",
            "difficulty": "Medium",
            "tags": json.dumps(["equilibrium", "haber", "thermodynamics"]),
            "year": "2025",
            "option_layout": "vertical",
            "explanation_text": "Le Chatelier principle favors forward reaction at high pressure (4 mol -> 2 mol) and lower temperature.",
            "diagram_svg": None,
            "diagram_url": None,
            "is_system": 1,
            "created_at": datetime.utcnow().isoformat(),
            "updated_at": datetime.utcnow().isoformat()
        },
        {
            "id": "q-chem-006",
            "question_number": 6,
            "question_type": "MCQ_SINGLE",
            "content": json.dumps([]),
            "raw_text": "The SI unit of the rate constant k for a second-order chemical reaction (n = 2) is:",
            "options": json.dumps([
                {"id": "opt-6a", "key": "a", "rawText": "\\text{L}\\cdot\\text{mol}^{-1}\\cdot\\text{s}^{-1}", "isCorrect": True, "content": []},
                {"id": "opt-6b", "key": "b", "rawText": "\\text{s}^{-1}", "isCorrect": False, "content": []},
                {"id": "opt-6c", "key": "c", "rawText": "\\text{mol}\\cdot\\text{L}^{-1}\\cdot\\text{s}^{-1}", "isCorrect": False, "content": []},
                {"id": "opt-6d", "key": "d", "rawText": "\\text{mol}^2\\cdot\\text{L}^{-2}\\cdot\\text{s}^{-1}", "isCorrect": False, "content": []}
            ]),
            "correct_answer": "a",
            "marks": 3,
            "negative_marks": 1,
            "subject": "Chemistry",
            "chapter": "Chemical Kinetics",
            "topic": "Order of Reaction",
            "difficulty": "Easy",
            "tags": json.dumps(["kinetics", "rate-constant", "units"]),
            "year": "2024",
            "option_layout": "grid_2x2",
            "explanation_text": "General formula for rate constant unit is (mol/L)^(1-n) s^(-1). For n=2: (mol/L)^(-1) s^(-1) = L mol^(-1) s^(-1).",
            "diagram_svg": None,
            "diagram_url": None,
            "is_system": 1,
            "created_at": datetime.utcnow().isoformat(),
            "updated_at": datetime.utcnow().isoformat()
        },
        {
            "id": "q-chem-007",
            "question_number": 7,
            "question_type": "MCQ_SINGLE",
            "content": json.dumps([]),
            "raw_text": "The IUPAC name of the coordination compound [Co(NH₃)₅(CO₃)]Cl is:",
            "options": json.dumps([
                {"id": "opt-7a", "key": "a", "rawText": "Pentaamminecarbonatocobalt(III) chloride", "isCorrect": True, "content": []},
                {"id": "opt-7b", "key": "b", "rawText": "Carbonatopentaamminecobalt(III) chloride", "isCorrect": False, "content": []},
                {"id": "opt-7c", "key": "c", "rawText": "Pentaamminecobalt(III) carbonate chloride", "isCorrect": False, "content": []},
                {"id": "opt-7d", "key": "d", "rawText": "Pentaamminecarbonatocobalt(II) chloride", "isCorrect": False, "content": []}
            ]),
            "correct_answer": "a",
            "marks": 3,
            "negative_marks": 1,
            "subject": "Chemistry",
            "chapter": "Coordination Compounds",
            "topic": "IUPAC Nomenclature",
            "difficulty": "Medium",
            "tags": json.dumps(["coordination", "inorganic", "iupac"]),
            "year": "2025",
            "option_layout": "vertical",
            "explanation_text": "Ligands in alphabetical order: ammine before carbonato. Oxidation state of Co is +3: Co + 5(0) + (-2) + (-1) = 0 => Co = +3.",
            "diagram_svg": None,
            "diagram_url": None,
            "is_system": 1,
            "created_at": datetime.utcnow().isoformat(),
            "updated_at": datetime.utcnow().isoformat()
        },
        {
            "id": "q-chem-008",
            "question_number": 8,
            "question_type": "MCQ_SINGLE",
            "content": json.dumps([]),
            "raw_text": "In the nitration of benzene using concentrated HNO₃ and concentrated H₂SO₄, the active electrophile is:",
            "options": json.dumps([
                {"id": "opt-8a", "key": "a", "rawText": "\\text{NO}_2^+ \\text{ (Nitronium ion)}", "isCorrect": True, "content": []},
                {"id": "opt-8b", "key": "b", "rawText": "\\text{NO}^+ \\text{ (Nitrosonium ion)}", "isCorrect": False, "content": []},
                {"id": "opt-8c", "key": "c", "rawText": "\\text{NO}_3^- \\text{ (Nitrate ion)}", "isCorrect": False, "content": []},
                {"id": "opt-8d", "key": "d", "rawText": "\\text{HNO}_2 \\text{ (Nitrous acid)}", "isCorrect": False, "content": []}
            ]),
            "correct_answer": "a",
            "marks": 3,
            "negative_marks": 1,
            "subject": "Chemistry",
            "chapter": "Hydrocarbons & Aromatics",
            "topic": "Electrophilic Aromatic Substitution",
            "difficulty": "Easy",
            "tags": json.dumps(["organic", "benzene", "nitration"]),
            "year": "2025",
            "option_layout": "grid_2x2",
            "explanation_text": "HNO3 + 2 H2SO4 -> NO2+ + H3O+ + 2 HSO4-.",
            "diagram_svg": None,
            "diagram_url": None,
            "is_system": 1,
            "created_at": datetime.utcnow().isoformat(),
            "updated_at": datetime.utcnow().isoformat()
        },
        {
            "id": "q-math-009",
            "question_number": 9,
            "question_type": "MCQ_SINGLE",
            "content": json.dumps([]),
            "raw_text": "Evaluate the definite integral: I = \\int_{0}^{\\pi/2} \\frac{\\sqrt{\\sin x}}{\\sqrt{\\sin x} + \\sqrt{\\cos x}}\\,dx",
            "options": json.dumps([
                {"id": "opt-9a", "key": "a", "rawText": "\\frac{\\pi}{4}", "isCorrect": True, "content": []},
                {"id": "opt-9b", "key": "b", "rawText": "\\frac{\\pi}{2}", "isCorrect": False, "content": []},
                {"id": "opt-9c", "key": "c", "rawText": "\\pi", "isCorrect": False, "content": []},
                {"id": "opt-9d", "key": "d", "rawText": "0", "isCorrect": False, "content": []}
            ]),
            "correct_answer": "a",
            "marks": 4,
            "negative_marks": 0,
            "subject": "Mathematics",
            "chapter": "Definite Integrals",
            "topic": "Properties of Integration",
            "difficulty": "Easy",
            "tags": json.dumps(["calculus", "integration"]),
            "year": "2024",
            "option_layout": "grid_2x2",
            "explanation_text": "Apply King's property: I + I = \\int_0^{\\pi/2} 1 dx = \\pi/2 => I = \\pi/4.",
            "diagram_svg": None,
            "diagram_url": None,
            "is_system": 1,
            "created_at": datetime.utcnow().isoformat(),
            "updated_at": datetime.utcnow().isoformat()
        },
        {
            "id": "q-math-010",
            "question_number": 10,
            "question_type": "MCQ_SINGLE",
            "content": json.dumps([]),
            "raw_text": "If \\vec{a} and \\vec{b} are two unit vectors inclined at an angle \\theta, then |\\vec{a} - \\vec{b}| is equal to:",
            "options": json.dumps([
                {"id": "opt-10a", "key": "a", "rawText": "2 \\sin\\left(\\frac{\\theta}{2}\\right)", "isCorrect": True, "content": []},
                {"id": "opt-10b", "key": "b", "rawText": "2 \\cos\\left(\\frac{\\theta}{2}\\right)", "isCorrect": False, "content": []},
                {"id": "opt-10c", "key": "c", "rawText": "\\sin\\theta", "isCorrect": False, "content": []},
                {"id": "opt-10d", "key": "d", "rawText": "\\cos\\theta", "isCorrect": False, "content": []}
            ]),
            "correct_answer": "a",
            "marks": 4,
            "negative_marks": 1,
            "subject": "Mathematics",
            "chapter": "Vector Algebra",
            "topic": "Dot and Cross Products",
            "difficulty": "Medium",
            "tags": json.dumps(["vectors", "trigonometry"]),
            "year": "2025",
            "option_layout": "grid_2x2",
            "explanation_text": "|a - b|^2 = |a|^2 + |b|^2 - 2(a.b) = 1 + 1 - 2 cos theta = 2(1 - cos theta) = 4 sin^2(theta/2) => |a - b| = 2 sin(theta/2).",
            "diagram_svg": None,
            "diagram_url": None,
            "is_system": 1,
            "created_at": datetime.utcnow().isoformat(),
            "updated_at": datetime.utcnow().isoformat()
        }
    ]

    for q in sample_questions:
        cursor.execute("""
        INSERT INTO question_bank (
            id, question_number, question_type, content, raw_text, options, correct_answer,
            marks, negative_marks, subject, chapter, topic, difficulty, tags, year,
            option_layout, explanation_text, diagram_svg, diagram_url, is_system, created_at, updated_at
        ) VALUES (
            :id, :question_number, :question_type, :content, :raw_text, :options, :correct_answer,
            :marks, :negative_marks, :subject, :chapter, :topic, :difficulty, :tags, :year,
            :option_layout, :explanation_text, :diagram_svg, :diagram_url, :is_system, :created_at, :updated_at
        )
        """, q)


def seed_templates(cursor):
    templates = [
        {
            "id": "tpl-cbse-12",
            "name": "CBSE Class 12 Board Standard",
            "description": "2-Column standard layout with Section A (MCQs), Section B (Short), Section C (Long)",
            "category": "CBSE",
            "settings": json.dumps({
                "pageSize": "A4",
                "columns": 2,
                "columnGap": 8,
                "columnDivider": True,
                "margins": {"top": 12, "bottom": 12, "left": 12, "right": 12},
                "defaultFont": "Inter",
                "defaultFontSize": 10.5,
                "showPageNumbers": True
            }),
            "default_metadata": json.dumps({
                "instituteName": "DELHI PUBLIC ACADEMY",
                "examName": "ALL INDIA SENIOR SCHOOL CERTIFICATE EXAMINATION 2026",
                "subject": "PHYSICS (THEORY)",
                "grade": "CLASS XII",
                "timeAllowedMinutes": 180,
                "maxMarks": 70,
                "generalInstructions": [
                    "All questions are compulsory. There are 33 questions in total.",
                    "Section A contains 16 MCQs carrying 1 mark each.",
                    "Section B contains 5 Short Answer questions of 2 marks each.",
                    "Use of log tables and mathematical tables is permitted."
                ]
            }),
            "default_sections": json.dumps([
                {"id": "sec-1", "defaultTitle": "SECTION A (MCQs)", "defaultInstructions": "Q.1 to Q.16 carry 1 mark each. Select the correct option.", "defaultMarks": 16},
                {"id": "sec-2", "defaultTitle": "SECTION B (Short Answer I)", "defaultInstructions": "Q.17 to Q.21 carry 2 marks each.", "defaultMarks": 10},
                {"id": "sec-3", "defaultTitle": "SECTION C (Long Answer)", "defaultInstructions": "Q.22 to Q.28 carry 3 marks each.", "defaultMarks": 21}
            ]),
            "created_at": datetime.utcnow().isoformat()
        },
        {
            "id": "tpl-jee-main",
            "name": "JEE Main / Advanced NTA Format",
            "description": "Standard engineering entrance exam format with Section 1 (Single Correct) and Section 2 (Numerical).",
            "category": "Competitive",
            "settings": json.dumps({
                "pageSize": "A4",
                "columns": 2,
                "columnGap": 8,
                "columnDivider": True,
                "margins": {"top": 10, "bottom": 10, "left": 10, "right": 10},
                "defaultFont": "Inter",
                "defaultFontSize": 10.0,
                "showPageNumbers": True
            }),
            "default_metadata": json.dumps({
                "instituteName": "NATIONAL TESTING AGENCY",
                "examName": "JOINT ENTRANCE EXAMINATION (MAIN) - 2026",
                "subject": "PHYSICS, CHEMISTRY & MATHEMATICS",
                "grade": "JEE ADVANCED / MAIN",
                "timeAllowedMinutes": 180,
                "maxMarks": 300,
                "generalInstructions": [
                    "This question paper contains 90 questions divided into 3 parts: Physics, Chemistry, and Mathematics.",
                    "Each part contains Section A (+4, -1) and Section B (+4, 0).",
                    "Calculators and electronic devices are strictly prohibited."
                ]
            }),
            "default_sections": json.dumps([
                {"id": "sec-jee-1", "defaultTitle": "SECTION I - PHYSICS (Single Choice)", "defaultInstructions": "+4 for correct, -1 for incorrect", "defaultMarks": 80},
                {"id": "sec-jee-2", "defaultTitle": "SECTION II - CHEMISTRY (Single Choice)", "defaultInstructions": "+4 for correct, -1 for incorrect", "defaultMarks": 80},
                {"id": "sec-jee-3", "defaultTitle": "SECTION III - MATHEMATICS (Single Choice)", "defaultInstructions": "+4 for correct, -1 for incorrect", "defaultMarks": 80}
            ]),
            "created_at": datetime.utcnow().isoformat()
        }
    ]

    for t in templates:
        cursor.execute("""
        INSERT INTO templates (id, name, description, category, settings, default_metadata, default_sections, created_at)
        VALUES (:id, :name, :description, :category, :settings, :default_metadata, :default_sections, :created_at)
        """, t)


# Initialize DB on startup
init_db()


# --------------------------------------------------------------------------
# API Routes
# --------------------------------------------------------------------------

@app.get("/api/health")
def health():
    return {
        "status": "ok",
        "name": "EduForge Python API",
        "version": "2.0.0",
        "engine": "FastAPI + SQLite",
        "timestamp": datetime.utcnow().isoformat()
    }


# --- Documents CRUD ---

@app.get("/api/documents")
def list_documents():
    conn = get_db()
    cursor = conn.cursor()
    cursor.execute("SELECT * FROM documents ORDER BY updated_at DESC")
    rows = cursor.fetchall()
    conn.close()
    
    docs = []
    for r in rows:
        docs.append({
            "id": r["id"],
            "title": r["title"],
            "templateId": r["template_id"],
            "settings": json.loads(r["settings"]),
            "metadata": json.loads(r["metadata"]),
            "sections": json.loads(r["sections"]),
            "createdAt": r["created_at"],
            "updatedAt": r["updated_at"]
        })
    return {"success": True, "data": docs}


@app.post("/api/documents")
def create_document(payload: Dict[str, Any]):
    doc_id = payload.get("id") or f"doc-{uuid.uuid4().hex[:8]}"
    now = datetime.utcnow().isoformat()
    
    title = payload.get("title", "Untitled Question Paper")
    template_id = payload.get("templateId")
    settings = json.dumps(payload.get("settings", {
        "pageSize": "A4",
        "columns": 2,
        "columnGap": 8,
        "columnDivider": True,
        "margins": {"top": 12, "bottom": 12, "left": 12, "right": 12},
        "defaultFont": "Inter",
        "defaultFontSize": 10.5,
        "showPageNumbers": True
    }))
    metadata = json.dumps(payload.get("metadata", {
        "instituteName": "MODEL PUBLIC SCHOOL",
        "examName": "ANNUAL EXAMINATION 2026",
        "subject": "Science & Mathematics",
        "timeAllowedMinutes": 180,
        "maxMarks": 100,
        "generalInstructions": ["All questions are compulsory."]
    }))
    sections = json.dumps(payload.get("sections", [
        {
            "id": f"sec-{uuid.uuid4().hex[:6]}",
            "title": "SECTION A",
            "instructions": "Answer all questions",
            "marks": 20,
            "blocks": []
        }
    ]))

    conn = get_db()
    cursor = conn.cursor()
    cursor.execute("""
    INSERT INTO documents (id, title, template_id, settings, metadata, sections, created_at, updated_at)
    VALUES (?, ?, ?, ?, ?, ?, ?, ?)
    """, (doc_id, title, template_id, settings, metadata, sections, now, now))
    conn.commit()
    conn.close()

    return {
        "id": doc_id,
        "title": title,
        "templateId": template_id,
        "settings": json.loads(settings),
        "metadata": json.loads(metadata),
        "sections": json.loads(sections),
        "createdAt": now,
        "updatedAt": now
    }


@app.get("/api/documents/{doc_id}")
def get_document(doc_id: str):
    conn = get_db()
    cursor = conn.cursor()
    cursor.execute("SELECT * FROM documents WHERE id = ?", (doc_id,))
    row = cursor.fetchone()
    conn.close()
    
    if not row:
        raise HTTPException(status_code=404, detail="Document not found")
    
    return {
        "id": row["id"],
        "title": row["title"],
        "templateId": row["template_id"],
        "settings": json.loads(row["settings"]),
        "metadata": json.loads(row["metadata"]),
        "sections": json.loads(row["sections"]),
        "createdAt": row["created_at"],
        "updatedAt": row["updated_at"]
    }


@app.put("/api/documents/{doc_id}")
def update_document(doc_id: str, payload: Dict[str, Any]):
    now = datetime.utcnow().isoformat()
    conn = get_db()
    cursor = conn.cursor()
    
    title = payload.get("title", "Untitled Document")
    template_id = payload.get("templateId")
    settings = json.dumps(payload.get("settings", {}))
    metadata = json.dumps(payload.get("metadata", {}))
    sections = json.dumps(payload.get("sections", []))

    cursor.execute("""
    UPDATE documents
    SET title = ?, template_id = ?, settings = ?, metadata = ?, sections = ?, updated_at = ?
    WHERE id = ?
    """, (title, template_id, settings, metadata, sections, now, doc_id))
    
    conn.commit()
    conn.close()
    return {"success": True, "updatedAt": now}


@app.delete("/api/documents/{doc_id}")
def delete_document(doc_id: str):
    conn = get_db()
    cursor = conn.cursor()
    cursor.execute("DELETE FROM documents WHERE id = ?", (doc_id,))
    conn.commit()
    conn.close()
    return {"success": True, "message": "Document deleted"}


def safe_json_loads(val: Any, default: Any = None) -> Any:
    if val is None:
        return default
    if isinstance(val, (dict, list)):
        return val
    if not isinstance(val, str):
        return default
    try:
        return json.loads(val)
    except Exception:
        try:
            import re
            sanitized = re.sub(r'\\(?!["\\/bfnrtu])', r'\\\\', val)
            return json.loads(sanitized)
        except Exception:
            return default if default is not None else val


# --- Question Bank API ---

@app.get("/api/question-bank")
def list_questions(
    subject: Optional[str] = Query(None),
    chapter: Optional[str] = Query(None),
    topic: Optional[str] = Query(None),
    difficulty: Optional[str] = Query(None),
    year: Optional[str] = Query(None),
    question_type: Optional[str] = Query(None),
    search: Optional[str] = Query(None),
    marks_min: Optional[float] = Query(None),
    marks_max: Optional[float] = Query(None),
):
    conn = get_db()
    cursor = conn.cursor()

    query = "SELECT * FROM question_bank WHERE 1=1"
    params = []

    if subject and subject != "all":
        query += " AND subject = ?"
        params.append(subject)
    if chapter:
        query += " AND chapter LIKE ?"
        params.append(f"%{chapter}%")
    if topic:
        query += " AND topic LIKE ?"
        params.append(f"%{topic}%")
    if difficulty and difficulty != "all":
        query += " AND difficulty = ?"
        params.append(difficulty)
    if year:
        query += " AND year = ?"
        params.append(year)
    if question_type:
        query += " AND question_type = ?"
        params.append(question_type)
    if search:
        query += " AND (raw_text LIKE ? OR chapter LIKE ? OR topic LIKE ?)"
        params.extend([f"%{search}%", f"%{search}%", f"%{search}%"])
    if marks_min is not None:
        query += " AND marks >= ?"
        params.append(marks_min)
    if marks_max is not None:
        query += " AND marks <= ?"
        params.append(marks_max)

    query += " ORDER BY question_number ASC, created_at DESC"
    cursor.execute(query, params)
    rows = cursor.fetchall()
    conn.close()

    questions = []
    for r in rows:
        questions.append({
            "id": r["id"],
            "questionNumber": r["question_number"],
            "questionType": r["question_type"],
            "content": safe_json_loads(r["content"], []),
            "rawText": r["raw_text"],
            "options": safe_json_loads(r["options"], []),
            "correctAnswer": r["correct_answer"],
            "marks": r["marks"],
            "negativeMarks": r["negative_marks"],
            "subject": r["subject"],
            "chapter": r["chapter"],
            "topic": r["topic"],
            "difficulty": r["difficulty"],
            "tags": safe_json_loads(r["tags"], []),
            "year": r["year"],
            "optionLayout": r["option_layout"],
            "explanationText": r["explanation_text"],
            "diagramSvg": r["diagram_svg"],
            "diagramUrl": r["diagram_url"],
            "isSystem": bool(r["is_system"]),
            "createdAt": r["created_at"],
            "updatedAt": r["updated_at"]
        })

    return {"success": True, "data": questions, "total": len(questions)}


@app.get("/api/question-bank/{q_id}")
def get_single_question(q_id: str):
    conn = get_db()
    cursor = conn.cursor()
    cursor.execute("SELECT * FROM question_bank WHERE id = ?", (q_id,))
    r = cursor.fetchone()
    conn.close()

    if not r:
        raise HTTPException(status_code=404, detail="Question not found")

    return {
        "id": r["id"],
        "questionNumber": r["question_number"],
        "questionType": r["question_type"],
        "content": safe_json_loads(r["content"], []),
        "rawText": r["raw_text"],
        "options": safe_json_loads(r["options"], []),
        "correctAnswer": r["correct_answer"],
        "marks": r["marks"],
        "negativeMarks": r["negative_marks"],
        "subject": r["subject"],
        "chapter": r["chapter"],
        "topic": r["topic"],
        "difficulty": r["difficulty"],
        "tags": safe_json_loads(r["tags"], []),
        "year": r["year"],
        "optionLayout": r["option_layout"],
        "explanationText": r["explanation_text"],
        "diagramSvg": r["diagram_svg"],
        "diagramUrl": r["diagram_url"],
        "isSystem": bool(r["is_system"]),
        "createdAt": r["created_at"],
        "updatedAt": r["updated_at"]
    }


@app.post("/api/question-bank")
def create_question(payload: Dict[str, Any]):
    q_id = payload.get("id") or f"q-{uuid.uuid4().hex[:8]}"
    now = datetime.utcnow().isoformat()
    
    conn = get_db()
    cursor = conn.cursor()
    cursor.execute("""
    INSERT INTO question_bank (
        id, question_number, question_type, content, raw_text, options, correct_answer,
        marks, negative_marks, subject, chapter, topic, difficulty, tags, year,
        option_layout, explanation_text, diagram_svg, diagram_url, is_system, created_at, updated_at
    ) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
    """, (
        q_id,
        payload.get("questionNumber", 1),
        payload.get("questionType", "MCQ_SINGLE"),
        json.dumps(payload.get("content", [])),
        payload.get("rawText", ""),
        json.dumps(payload.get("options", [])),
        payload.get("correctAnswer"),
        payload.get("marks", 1),
        payload.get("negativeMarks", 0),
        payload.get("subject", "Physics"),
        payload.get("chapter", "General"),
        payload.get("topic", "General"),
        payload.get("difficulty", "Medium"),
        json.dumps(payload.get("tags", [])),
        payload.get("year", "2026"),
        payload.get("optionLayout", "grid_2x2"),
        payload.get("explanationText", ""),
        payload.get("diagramSvg"),
        payload.get("diagramUrl"),
        0,
        now,
        now
    ))
    conn.commit()
    conn.close()

    return {"success": True, "id": q_id, "createdAt": now}


@app.delete("/api/question-bank/{q_id}")
def delete_question(q_id: str):
    conn = get_db()
    cursor = conn.cursor()
    cursor.execute("DELETE FROM question_bank WHERE id = ?", (q_id,))
    conn.commit()
    conn.close()
    return {"success": True, "message": "Question deleted"}


# --- Templates API ---

@app.get("/api/templates")
def list_templates():
    conn = get_db()
    cursor = conn.cursor()
    cursor.execute("SELECT * FROM templates ORDER BY created_at ASC")
    rows = cursor.fetchall()
    conn.close()

    templates = []
    for r in rows:
        templates.append({
            "id": r["id"],
            "name": r["name"],
            "description": r["description"],
            "category": r["category"],
            "settings": json.loads(r["settings"]),
            "defaultMetadata": json.loads(r["default_metadata"]),
            "defaultSections": json.loads(r["default_sections"]),
            "createdAt": r["created_at"]
        })
    return {"success": True, "data": templates}


# --- Science Catalogs ---

RESOURCES_DIR = os.path.abspath(os.path.join(os.path.dirname(__file__), "..", "resources"))

def load_resource_json(rel_path: str, default: Any) -> Any:
    full_path = os.path.join(RESOURCES_DIR, rel_path)
    if os.path.exists(full_path):
        try:
            with open(full_path, "r", encoding="utf-8") as f:
                return json.load(f)
        except Exception:
            pass
    return default


@app.get("/api/physics")
def get_physics_catalog():
    return {
        "success": True,
        "chapters": [
            {
                "id": "ch-kinematics",
                "title": "Kinematics & Projectiles",
                "formulas": [
                    {"name": "Velocity-Time", "latex": "v = u + at"},
                    {"name": "Displacement-Time", "latex": "s = ut + \\frac{1}{2}at^2"},
                    {"name": "Velocity-Displacement", "latex": "v^2 = u^2 + 2as"},
                    {"name": "Max Height Projectile", "latex": "H = \\frac{u^2\\sin^2\\theta}{2g}"},
                    {"name": "Time of Flight", "latex": "T = \\frac{2u\\sin\\theta}{g}"},
                    {"name": "Horizontal Range", "latex": "R = \\frac{u^2\\sin(2\\theta)}{g}"}
                ]
            },
            {
                "id": "ch-electromagnetism",
                "title": "Electromagnetism & Maxwell",
                "formulas": [
                    {"name": "Coulomb Law", "latex": "F = \\frac{1}{4\\pi\\varepsilon_0}\\frac{|q_1 q_2|}{r^2}"},
                    {"name": "Electric Field Point Charge", "latex": "\\vec{E} = \\frac{q}{4\\pi\\varepsilon_0 r^2}\\hat{r}"},
                    {"name": "Lorentz Force", "latex": "\\vec{F} = q(\\vec{E} + \\vec{v}\\times\\vec{B})"},
                    {"name": "Faraday Induction", "latex": "\\mathcal{E} = -\\frac{d\\Phi_B}{dt}"},
                    {"name": "Biot-Savart Law", "latex": "d\\vec{B} = \\frac{\\mu_0}{4\\pi}\\frac{I d\\vec{l}\\times\\hat{r}}{r^2}"}
                ]
            }
        ]
    }


@app.get("/api/physics/chapters")
def get_physics_chapters():
    return load_resource_json("physics/chapters.json", [])


@app.get("/api/chemistry")
def get_chemistry_catalog():
    return {
        "success": True,
        "reactions": [
            {"title": "Haber Ammonia", "latex": "\\text{N}_2\\text{(g)} + 3\\text{H}_2\\text{(g)} \\rightleftharpoons 2\\text{NH}_3\\text{(g)}"},
            {"title": "Combustion of Methane", "latex": "\\text{CH}_4 + 2\\text{O}_2 \\rightarrow \\text{CO}_2 + 2\\text{H}_2\\text{O}"},
            {"title": "Nernst Equation", "latex": "E_{\\text{cell}} = E^\\circ_{\\text{cell}} - \\frac{0.0591}{n}\\log_{10}Q"},
            {"title": "Arrhenius Equation", "latex": "k = A e^{-\\frac{E_a}{RT}}"}
        ]
    }


@app.get("/api/chemistry/elements")
def get_chemistry_elements():
    return load_resource_json("chemistry/elements.json", [])


@app.get("/api/chemistry/notations")
def get_chemistry_notations():
    return load_resource_json("chemistry/notations.json", [])


@app.get("/api/units")
def get_units():
    data = load_resource_json("units/units.json", {"units": [], "prefixes": []})
    if isinstance(data, dict):
        return data.get("units", [])
    return data if isinstance(data, list) else []


@app.get("/api/units/prefixes")
def get_prefixes():
    data = load_resource_json("units/units.json", {"units": [], "prefixes": []})
    if isinstance(data, dict):
        return data.get("prefixes", [])
    return []


@app.get("/api/constants")
def get_constants():
    return load_resource_json("constants/constants.json", [])


@app.get("/api/symbols")
def get_symbols_catalog():
    return load_resource_json("symbols/symbols.json", {
        "success": True,
        "greek": ["\\alpha", "\\beta", "\\gamma", "\\delta", "\\theta", "\\lambda", "\\mu", "\\pi", "\\sigma", "\\omega", "\\Delta", "\\Sigma", "\\Omega"],
        "operators": ["\\pm", "\\times", "\\div", "\\cdot", "\\approx", "\\le", "\\ge", "\\neq", "\\propto", "\\int", "\\sum", "\\prod"]
    })


# --- Export DOCX & PDF ---

@app.post("/api/export/docx")
def export_docx(doc: Dict[str, Any]):
    from docx import Document
    from docx.shared import Inches, Pt
    import io

    document = Document()
    
    title = doc.get("title", "Question Paper")
    heading = document.add_heading(title, level=0)
    heading.alignment = 1

    metadata = doc.get("metadata", {})
    inst_name = metadata.get("instituteName") or metadata.get("institutionName")
    if inst_name:
        p = document.add_paragraph()
        p.alignment = 1
        run = p.add_run(inst_name.upper())
        run.bold = True
        run.font.size = Pt(14)

    if metadata.get("examName"):
        p = document.add_paragraph()
        p.alignment = 1
        p.add_run(metadata["examName"]).bold = True

    # Info table (Time, Max Marks)
    info_table = document.add_table(rows=1, cols=2)
    info_table.autofit = True
    r = info_table.rows[0]
    r.cells[0].text = f"Time Allowed: {metadata.get('timeAllowedMinutes', 180)} Minutes"
    r.cells[1].text = f"Maximum Marks: {metadata.get('maxMarks', 100)}"
    r.cells[1].paragraphs[0].alignment = 2

    document.add_paragraph()

    # Sections & Blocks
    sections = doc.get("sections", [])
    for sec in sections:
        sec_title = document.add_heading(sec.get("title", "SECTION"), level=2)
        sec_title.alignment = 1
        if sec.get("instructions"):
            p_inst = document.add_paragraph(sec["instructions"])
            p_inst.italic = True
            p_inst.alignment = 1

        for block in sec.get("blocks", []):
            btype = block.get("type")
            if btype == "question":
                q = block.get("question", {})
                q_num = q.get("questionNumber", "")
                q_text = q.get("rawText", "")
                marks = q.get("marks", 1)

                p = document.add_paragraph()
                p.add_run(f"{q_num}. ").bold = True
                p.add_run(q_text)
                p.add_run(f"  [{marks} Marks]").bold = True

                options = q.get("options", [])
                for opt in options:
                    opt_p = document.add_paragraph()
                    opt_p.paragraph_format.left_indent = Inches(0.5)
                    key = opt.get("key", "a").upper()
                    raw = opt.get("rawText", "")
                    opt_p.add_run(f"({key}) {raw}")

            elif btype == "paragraph":
                runs = block.get("runs", [])
                p = document.add_paragraph()
                for r_item in runs:
                    p.add_run(r_item.get("text", ""))

            elif btype == "equation":
                latex = block.get("rawLatex", "")
                p = document.add_paragraph()
                p.alignment = 1
                p.add_run(f"[Formula: {latex}]").italic = True

    buf = io.BytesIO()
    document.save(buf)
    buf.seek(0)
    
    return Response(
        content=buf.getvalue(),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{title}.docx"'}
    )


@app.post("/api/export/pdf-html")
def export_pdf_html(doc: Dict[str, Any]):
    title = doc.get("title", "Question Paper")
    metadata = doc.get("metadata", {})
    sections = doc.get("sections", [])
    inst_name = metadata.get("instituteName") or metadata.get("institutionName") or "EXAMINATION ACADEMY"
    
    html = f"""<!DOCTYPE html>
<html>
<head>
    <meta charset="utf-8"/>
    <title>{title}</title>
    <link rel="stylesheet" href="https://cdn.jsdelivr.net/npm/katex@0.16.9/dist/katex.min.css">
    <script src="https://cdn.jsdelivr.net/npm/katex@0.16.9/dist/katex.min.js"></script>
    <script src="https://cdn.jsdelivr.net/npm/katex@0.16.9/dist/contrib/auto-render.min.js"></script>
    <style>
        @page {{ size: A4; margin: 12mm; }}
        body {{ font-family: 'Inter', sans-serif; color: #0f172a; line-height: 1.4; margin: 0; padding: 12mm; }}
        .header {{ text-align: center; border-bottom: 2px solid #0f172a; padding-bottom: 8px; margin-bottom: 12px; }}
        .inst {{ font-size: 16pt; font-weight: 900; text-transform: uppercase; margin: 0; }}
        .exam {{ font-size: 12pt; font-weight: 700; margin: 4px 0; }}
        .meta-row {{ display: flex; justify-content: space-between; font-size: 10pt; font-weight: 600; margin-top: 6px; }}
        .section-title {{ text-align: center; font-size: 11pt; font-weight: 900; text-transform: uppercase; border-bottom: 1px solid #94a3b8; padding: 4px 0; margin-top: 16px; }}
        .question {{ margin: 10px 0; font-size: 10.5pt; }}
        .options-grid {{ display: grid; grid-template-columns: 1fr 1fr; gap: 6px 20px; margin: 6px 0 6px 16px; font-size: 10pt; }}
        .diagram {{ text-align: center; margin: 8px 0; }}
    </style>
</head>
<body>
    <div class="header">
        <div class="inst">{inst_name}</div>
        <div class="exam">{metadata.get('examName', 'QUESTION PAPER')}</div>
        <div class="meta-row">
            <span>Subject: {metadata.get('subject', 'General')}</span>
            <span>Time: {metadata.get('timeAllowedMinutes', 180)} Min | Max Marks: {metadata.get('maxMarks', 100)}</span>
        </div>
    </div>
"""

    for s in sections:
        html += f"<div class='section-title'>{s.get('title', 'SECTION')}</div>"
        if s.get("instructions"):
            html += f"<div style='text-align:center; font-style:italic; font-size:9pt; color:#64748b;'>{s.get('instructions')}</div>"
        
        for b in s.get("blocks", []):
            if b.get("type") == "question":
                q = b.get("question", {})
                q_num = q.get("questionNumber", "")
                q_text = q.get("rawText", "")
                marks = q.get("marks", 1)
                html += f"<div class='question'><strong>{q_num}.</strong> {q_text} <span style='float:right; font-size:9pt; font-weight:bold;'>[{marks}M]</span></div>"
                
                # Render diagram if present
                if q.get("diagramSvg"):
                    html += f"<div class='diagram'>{q['diagramSvg']}</div>"
                elif q.get("diagramUrl"):
                    html += f"<div class='diagram'><img src='{q['diagramUrl']}' style='max-height:160px;'/></div>"

                opts = q.get("options", [])
                if opts:
                    html += "<div class='options-grid'>"
                    for opt in opts:
                        key = opt.get("key", "a").upper()
                        raw = opt.get("rawText", "")
                        html += f"<div><strong>({key})</strong> {raw}</div>"
                    html += "</div>"

    html += """
    <script>
        document.addEventListener("DOMContentLoaded", function() {
            renderMathInElement(document.body, {
                delimiters: [
                    {left: "$$", right: "$$", display: true},
                    {left: "$", right: "$", display: false},
                    {left: "\\\\(", right: "\\\\)", display: false},
                    {left: "\\\\[", right: "\\\\]", display: true}
                ]
            });
        });
    </script>
</body>
</html>"""
    
    return Response(content=html, media_type="text/html")


@app.post("/api/assets/upload")
@app.post("/api/upload")
@app.post("/api/upload/image")
async def upload_asset(file: Optional[UploadFile] = None, image: Optional[UploadFile] = None):
    uploaded = file or image
    if not uploaded:
        raise HTTPException(status_code=400, detail="No file provided")
    
    ext = os.path.splitext(uploaded.filename)[1] or ".png"
    safe_name = f"{uuid.uuid4().hex[:10]}{ext}"
    filepath = os.path.join(UPLOADS_DIR, safe_name)
    
    with open(filepath, "wb") as f:
        content = await uploaded.read()
        f.write(content)
        
    url = f"/api/assets/{safe_name}"
    return {
        "success": True,
        "originalName": uploaded.filename,
        "filename": safe_name,
        "url": url,
        "data": {
            "id": f"asset-{uuid.uuid4().hex[:6]}",
            "url": url,
            "originalName": uploaded.filename
        }
    }


@app.get("/api/assets/{filename}")
@app.get("/api/uploads/{filename}")
async def get_uploaded_asset(filename: str):
    filepath = os.path.join(UPLOADS_DIR, filename)
    if not os.path.exists(filepath):
        raise HTTPException(status_code=404, detail="Asset not found")
    return FileResponse(filepath)


# Settings API
@app.get("/api/settings")
def get_app_settings():
    return {
        "success": True,
        "data": {
            "defaultFont": "Calibri",
            "defaultFontSize": 10.5,
            "defaultPaperSize": "A4",
            "defaultMargins": {"top": 12, "bottom": 12, "left": 12, "right": 12},
            "defaultQuestionStyle": "number",
            "defaultOptionStyle": "grid_2x2",
            "defaultEquationSize": 16,
            "autosaveIntervalMs": 2000,
            "theme": "dark",
            "exportSettings": {"pdfDpi": 300, "embedFonts": True, "showPageNumbers": True},
            "backupSettings": {"autoBackupDaily": True, "maxBackupsToKeep": 10}
        }
    }


@app.put("/api/settings")
def update_app_settings(payload: Dict[str, Any]):
    return {"success": True, "data": payload}


if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=3001)
